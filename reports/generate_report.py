"""Génère reports/rapport.docx à partir du contenu structuré ci-dessous.

Mise en forme imposée : noir et blanc uniquement (aucune autre couleur),
police Tahoma, taille normale 12. Contient un sommaire, une table des
matières (champ Word, à mettre à jour à l'ouverture), un glossaire, et les
sections imposées par le cours (cf. reports/README.md).

Usage : python reports/generate_report.py
"""

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent
FIGURES_DIR = REPO_ROOT / "reports" / "figures"
OUTPUT_PATH = REPO_ROOT / "reports" / "rapport.docx"

BLACK = RGBColor(0, 0, 0)
FONT = "Tahoma"
BODY_SIZE = Pt(12)

figure_counter = 0
table_counter = 0


# ---------------------------------------------------------------------------
# Mise en forme
# ---------------------------------------------------------------------------

def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(8)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)

    heading_sizes = {
        "Title": Pt(22),
        "Heading 1": Pt(16),
        "Heading 2": Pt(14),
        "Heading 3": Pt(12),
    }
    for style_name, size in heading_sizes.items():
        style = doc.styles[style_name]
        style.font.name = FONT
        style.font.size = size
        style.font.color.rgb = BLACK
        style.font.bold = True

    # Le style Hyperlink (utilisé par Word pour les entrées de TOC) n'existe
    # pas encore dans un document python-docx vierge : si on ne le définit
    # pas nous-mêmes, Word le crée à la volée avec son bleu souligné par
    # défaut dès la mise à jour du champ TOC, ce qui violerait la contrainte
    # "aucune autre couleur". On le déclare donc explicitement en noir, sans
    # soulignement.
    if "Hyperlink" in doc.styles:
        hyperlink = doc.styles["Hyperlink"]
    else:
        hyperlink = doc.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
    hyperlink.font.name = FONT
    hyperlink.font.size = BODY_SIZE
    hyperlink.font.color.rgb = BLACK
    hyperlink.font.underline = False

    # Champs mis à jour automatiquement à l'ouverture (sommaire/TOC), pour
    # que les numéros de page apparaissent sans action manuelle dans Word.
    settings_element = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings_element.append(update_fields)


def add_toc_field(doc: Document, levels: str = "1-3") -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f'TOC \\o "{levels}" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Table des matières (clic droit → Mettre à jour les champs)"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_separate)
    r.append(placeholder)
    r.append(fld_end)


def add_figure(doc: Document, filename: str, caption: str, width_cm: float = 14.0) -> None:
    global figure_counter
    figure_counter += 1
    doc.add_picture(str(FIGURES_DIR / filename), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"Figure {figure_counter} — {caption}")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    global table_counter
    table_counter += 1
    cap = doc.add_paragraph()
    run = cap.add_run(f"Tableau {table_counter} — {caption}")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = FONT
        r.font.size = BODY_SIZE
        r.font.color.rgb = BLACK
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = FONT
            r.font.size = BODY_SIZE
            r.font.color.rgb = BLACK
    doc.add_paragraph()


def add_glossary_entry(doc: Document, term: str, definition: str) -> None:
    p = doc.add_paragraph()
    run_term = p.add_run(f"{term} — ")
    run_term.bold = True
    p.add_run(definition)


# ---------------------------------------------------------------------------
# Contenu
# ---------------------------------------------------------------------------

def build_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Modèle de diffusion (DDPM) entraîné from scratch\nvs. GAN (DCGAN)")
    r.font.name = FONT
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = BLACK

    for _ in range(3):
        doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(
        "Comparaison rigoureuse de deux familles de modèles génératifs sur Fashion-MNIST :\n"
        "qualité de génération, diversité, stabilité d'entraînement, coût de calcul."
    )
    r.font.name = FONT
    r.font.size = Pt(13)
    r.font.color.rgb = BLACK

    for _ in range(4):
        doc.add_paragraph()

    infos = doc.add_paragraph()
    infos.alignment = WD_ALIGN_PARAGRAPH.CENTER
    infos.add_run(
        "Projet annuel — cours Projets AI & Big Data (TCHAYE-KONDI Jude, Ph.D.), Master\n"
        "Sujet n°3\n\n"
        "Équipe :\n"
        "KOYE Leleda Ma Belle — Data / Experiment Engineer\n"
        "AHLI Kossi Sitsofe Pédro — Model / Research Engineer (principal), appui Reporting / Backend\n"
        "KONTEVI Akossiwa Anne — Reporting / Backend Developer\n\n"
        "2026-08-06"
    )
    doc.add_page_break()


def build_sommaire(doc: Document) -> None:
    doc.add_heading("Sommaire", level=1)
    entries = [
        "1. Introduction",
        "2. Travaux liés",
        "3. Données",
        "4. Méthode",
        "5. Étude d'ablation",
        "6. Résultats",
        "7. Déploiement",
        "8. Discussion et limites",
        "9. Répartition du travail",
        "Références",
        "Glossaire",
    ]
    for entry in entries:
        doc.add_paragraph(entry, style="List Bullet")
    doc.add_page_break()


def build_toc(doc: Document) -> None:
    doc.add_heading("Table des matières", level=1)
    note = doc.add_paragraph()
    run = note.add_run(
        "(Champ Word — clic droit sur la table ci-dessous puis « Mettre à jour les champs » "
        "si les numéros de page ne s'affichent pas automatiquement à l'ouverture.)"
    )
    run.italic = True
    run.font.size = Pt(10)
    add_toc_field(doc)
    doc.add_page_break()


def build_glossary(doc: Document) -> None:
    doc.add_heading("Glossaire", level=1)
    entries = [
        ("DDPM (Denoising Diffusion Probabilistic Model)", "Modèle génératif qui apprend à débruiter progressivement une image partant d'un bruit gaussien pur, par inversion d'un processus de diffusion direct (Ho et al., 2020)."),
        ("GAN (Generative Adversarial Network)", "Modèle génératif composé de deux réseaux en compétition : un générateur qui produit des images, un discriminateur qui tente de distinguer les images réelles des images générées (Goodfellow et al., 2014)."),
        ("DCGAN", "Variante convolutive du GAN, avec des recommandations d'architecture stabilisant l'entraînement (BatchNorm, LeakyReLU, pas de couches entièrement connectées) (Radford et al., 2015)."),
        ("U-Net", "Architecture de réseau de neurones en forme de U, avec des connexions résiduelles (skip connections) entre les couches d'encodage et de décodage ; utilisée ici pour prédire le bruit à débruiter à chaque pas de diffusion."),
        ("Bruitage / débruitage", "Processus direct (bruitage) qui transforme progressivement une image en bruit gaussien ; processus inverse (débruitage), appris par le U-Net, qui reconstruit une image à partir du bruit."),
        ("Timestep (pas de diffusion, t)", "Étape du processus de diffusion, indexée de 0 (image originale) à T (bruit pur). Le nombre total de pas T est un hyperparamètre clé (cf. étude d'ablation)."),
        ("Schedule de bruit (β_t)", "Séquence de coefficients qui contrôle la quantité de bruit ajoutée à chaque pas de diffusion (schedule linéaire ou cosine)."),
        ("Seed", "Valeur d'initialisation du générateur de nombres aléatoires, fixée pour garantir la reproductibilité des résultats."),
        ("Checkpoint", "Sauvegarde des poids d'un modèle à un instant donné de l'entraînement, permettant de reprendre ou réutiliser le modèle sans le ré-entraîner."),
        ("Epoch / step", "Un step correspond à une mise à jour des poids sur un batch d'images ; une epoch correspond à un passage complet sur le jeu d'entraînement."),
        ("Mode collapse", "Défaut d'entraînement d'un GAN où le générateur produit un ensemble d'images peu varié, ne couvrant qu'une partie de la diversité réelle des données."),
        ("Discriminateur / Générateur", "Dans un GAN, le discriminateur classe une image comme réelle ou générée ; le générateur produit des images à partir d'un vecteur de bruit latent, en tentant de tromper le discriminateur."),
        ("FID (Fréchet Inception Distance)", "Métrique de qualité/diversité d'images générées, mesurant la distance statistique entre les distributions de caractéristiques (extraites par un réseau Inception) des images réelles et générées. Plus la valeur est basse, plus les images générées ressemblent aux images réelles."),
        ("Variance intra-batch", "Variance des pixels calculée entre plusieurs images générées ensemble ; utilisée ici comme indicateur simple de diversité (une variance proche de zéro indiquerait un mode collapse)."),
        ("Batch size", "Nombre d'images traitées simultanément à chaque step d'entraînement."),
        ("API REST", "Interface de programmation exposant des fonctionnalités (ici, la génération d'images) via des requêtes HTTP standard."),
        ("Docker / conteneur", "Technologie permettant d'empaqueter une application et ses dépendances dans un environnement isolé et reproductible (conteneur), exécutable de façon identique sur toute machine disposant de Docker."),
        ("Volume Docker", "Mécanisme permettant de partager des fichiers entre la machine hôte et un conteneur, sans les intégrer à l'image Docker elle-même (utilisé ici pour les checkpoints entraînés)."),
    ]
    for term, definition in entries:
        add_glossary_entry(doc, term, definition)
    doc.add_page_break()


def build_introduction(doc: Document) -> None:
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Ce projet s'inscrit dans le cadre du cours « Projets AI & Big Data » (TCHAYE-KONDI Jude, Ph.D.), "
        "sujet n°3 : comparer, de façon rigoureuse, deux familles de modèles génératifs d'images entraînés "
        "sur le même dataset et avec un budget de calcul comparable — un modèle de diffusion débruitant "
        "(DDPM) implémenté from scratch, et un GAN convolutif (DCGAN)."
    )
    doc.add_paragraph(
        "La problématique centrale est double : (1) implémenter, sans bibliothèque de diffusion préexistante, "
        "l'ensemble du pipeline DDPM — processus de diffusion direct, U-Net de débruitage, processus inverse "
        "d'échantillonnage — et un DCGAN entraîné dans les mêmes conditions ; (2) comparer objectivement les "
        "deux approches selon plusieurs axes : qualité de génération (FID), diversité des échantillons, "
        "stabilité d'entraînement, et coût de calcul (temps d'entraînement et de génération)."
    )
    doc.add_paragraph(
        "Contribution du projet : un pipeline de données reproductible pour Fashion-MNIST ; une implémentation "
        "from scratch du processus de diffusion (forward et inverse) et d'un U-Net compact adapté à un budget "
        "de calcul CPU contraint (absence de GPU) ; un DCGAN entraîné dans des conditions comparables ; une "
        "étude d'ablation sur le nombre de pas de diffusion, révélant un compromis qualité/coût non trivial ; "
        "une comparaison chiffrée DDPM vs GAN (FID, diversité, temps de génération) ; et une démonstration "
        "déployée (API FastAPI + application Streamlit, conteneurisées avec Docker)."
    )
    doc.add_paragraph(
        "L'ensemble des décisions, expériences et résultats intermédiaires est journalisé de façon détaillée "
        "dans HISTORY.md du dépôt de code, qui constitue la source primaire des chiffres cités dans ce rapport."
    )


def build_related_work(doc: Document) -> None:
    doc.add_heading("2. Travaux liés", level=1)
    doc.add_paragraph(
        "Les GAN (Generative Adversarial Networks) ont été introduits par Goodfellow et al. (2014) comme un "
        "jeu à somme nulle entre un générateur et un discriminateur, entraînés simultanément. Cette formulation "
        "adversariale produit des échantillons de haute qualité en une seule passe, mais est réputée instable "
        "à entraîner (oscillations, mode collapse). Radford et al. (2015) ont proposé DCGAN, une variante "
        "convolutive avec des recommandations d'architecture (BatchNorm, LeakyReLU, absence de couches "
        "entièrement connectées) qui stabilisent significativement l'entraînement ; c'est cette architecture "
        "qui sert de base à notre implémentation GAN (cf. section Méthode)."
    )
    doc.add_paragraph(
        "Les modèles de diffusion débruitants (DDPM) ont été popularisés par Ho et al. (2020), qui formulent "
        "la génération comme l'inversion d'un processus de diffusion direct ajoutant progressivement du bruit "
        "gaussien à une image. Leur contribution clé est une formule fermée pour échantillonner directement "
        "x_t à partir de x_0 sans boucle, et une loss simplifiée équivalente à une régression du bruit ajouté "
        "(MSE), rendant l'entraînement d'un DDPM comparable en simplicité à celui d'un réseau de régression "
        "standard. Nichol & Dhariwal (2021) ont ensuite proposé des améliorations, dont un schedule de bruit "
        "« cosine » plutôt que linéaire ; ce schedule est implémenté dans notre code (cf. "
        "src/models/ddpm/diffusion.py) mais le schedule linéaire d'origine a été retenu pour la baseline, par "
        "cohérence avec Ho et al. (2020)."
    )
    doc.add_paragraph(
        "Sur le plan des artefacts visuels, Odena et al. (2016) documentent les artefacts en damier "
        "(« checkerboard artifacts ») produits par les couches de convolution transposée, couramment utilisées "
        "dans les générateurs de GAN comme dans le nôtre — un phénomène observé empiriquement sur nos "
        "échantillons GAN (cf. section Résultats)."
    )


def build_data(doc: Document) -> None:
    doc.add_heading("3. Données", level=1)

    doc.add_heading("3.1 Choix du dataset", level=2)
    doc.add_paragraph(
        "Le choix du dataset s'est porté sur Fashion-MNIST plutôt que CIFAR-10 downscalé, après comparaison "
        "empirique des deux candidats (téléchargement et chargement testés via torchvision.datasets)."
    )
    add_table(
        doc,
        "Comparaison des datasets candidats",
        ["Dataset", "Train", "Test", "Taille image", "Classes", "Poids disque", "Temps de chargement"],
        [
            ["Fashion-MNIST", "60 000", "10 000", "1×28×28", "10", "81,85 Mo", "33,9 s"],
            ["CIFAR-10", "50 000", "10 000", "3×32×32", "10", "177,59 Mo", "1413,6 s (~23 min)"],
        ],
    )
    doc.add_paragraph(
        "Décision : Fashion-MNIST. Justification principale : absence de GPU sur la machine de développement "
        "(CPU uniquement) et budget de calcul limité pour le rôle Model (~14 h estimées pour DDPM + GAN + "
        "étude d'ablation). Fashion-MNIST représente un volume de calcul par image environ 3,92 fois moindre "
        "que CIFAR-10 (niveaux de gris 28×28 contre RGB 32×32), tout en restant suffisamment complexe "
        "visuellement (10 classes de vêtements, certaines proches visuellement) pour une comparaison DDPM vs "
        "GAN pertinente."
    )

    doc.add_heading("3.2 Analyse exploratoire (EDA)", level=2)
    doc.add_paragraph(
        "L'analyse exploratoire (notebooks/01_eda_dataset.ipynb) n'a révélé aucune anomalie : 70 000 images "
        "en niveaux de gris 28×28 (uint8, pixels dans [0, 255]), aucune valeur manquante, aucune dimension "
        "incohérente, aucun doublon exact (hash MD5 sur les 60 000 images d'entraînement et 10 000 de test). "
        "Les classes sont parfaitement équilibrées : 6 000 images par classe en entraînement, 1 000 en test."
    )
    add_figure(doc, "eda_classes.png", "Distribution des classes — Fashion-MNIST (train/test)")
    add_figure(doc, "eda_samples.png", "Échantillons Fashion-MNIST (6 images par classe)")
    doc.add_paragraph(
        "Statistiques de pixels (train, échelle [0, 1]) : moyenne = 0,286, écart-type = 0,353 — base retenue "
        "pour la normalisation du pipeline de données."
    )

    doc.add_heading("3.3 Pipeline de chargement et de prétraitement", level=2)
    doc.add_paragraph(
        "Le pipeline (src/data/dataset.py, configs/data.yaml) applique : un padding de 28×28 vers 32×32 "
        "(pour permettre des divisions entières par deux successives dans le U-Net : 32 → 16 → 8), une "
        "normalisation vers [-1, 1] (adaptée à la diffusion), et un split validation de 10 % prélevé sur le "
        "train set officiel (seed fixé à 42, via un torch.Generator dédié) — le test set officiel restant "
        "réservé à l'évaluation finale. Le batch size retenu est 128. La reproductibilité du split "
        "train/validation a été vérifiée par test unitaire (mêmes indices entre deux appels avec le même seed)."
    )


def build_method(doc: Document) -> None:
    doc.add_heading("4. Méthode", level=1)

    doc.add_heading("4.1 DDPM — processus de diffusion et U-Net", level=2)
    doc.add_paragraph(
        "Le processus de diffusion direct est implémenté from scratch (src/models/ddpm/diffusion.py) via la "
        "formule fermée q(x_t | x_0) = √ᾱ_t · x_0 + √(1-ᾱ_t) · ε, permettant de bruiter une image à un pas t "
        "arbitraire sans boucle sur les pas intermédiaires. Deux schedules de bruit β_t sont implémentés "
        "(linéaire et cosine) ; la baseline retient le schedule linéaire (β_start = 1e-4, β_end = 0,02, "
        "T = 1000 pas), conformément à Ho et al. (2020)."
    )
    add_figure(doc, "forward_noising.png", "Bruitage progressif q(x_t | x_0) sur 4 images (t = 0 à 999)")
    doc.add_paragraph(
        "Le U-Net de débruitage (src/models/ddpm/unet.py) prédit le bruit ajouté à x_t sachant t. Architecture : "
        "embedding sinusoïdal du pas de temps, blocs résiduels (GroupNorm + SiLU + convolution, injection de "
        "l'embedding temporel), downsampling/upsampling avec connexions résiduelles (skip connections) sur "
        "3 résolutions (32 → 16 → 8). Choix délibéré : aucune couche d'attention, pour rester entraînable sur "
        "CPU (contrairement à l'architecture originale de Ho et al., 2020, qui inclut de l'auto-attention à "
        "16×16). Le modèle compte environ 3,54 millions de paramètres. L'entraînement minimise une loss MSE "
        "entre le bruit réel et le bruit prédit ; l'échantillonnage utilise un processus inverse ancestral "
        "(x_T, bruit gaussien pur, jusqu'à x_0, sur T pas)."
    )

    doc.add_heading("4.2 GAN — architecture DCGAN", level=2)
    doc.add_paragraph(
        "Le générateur (src/models/gan/generator.py) transforme un vecteur latent (dimension 100) en image "
        "32×32×1 via 4 blocs de convolution transposée + BatchNorm + ReLU, sortie Tanh (cohérente avec la "
        "normalisation [-1, 1] du pipeline). Le discriminateur (src/models/gan/discriminator.py) est le miroir "
        "convolutif (LeakyReLU, pas de BatchNorm sur la première couche, conformément à Radford et al., 2015), "
        "sortie logit brut compatible avec une loss BCE avec logits. Les deux réseaux sont initialisés avec des "
        "poids tirés de N(0, 0,02). L'entraînement alterne une mise à jour du discriminateur puis du générateur "
        "à chaque step, avec des optimiseurs Adam séparés (lr = 2e-4, β1 = 0,5, β2 = 0,999)."
    )

    doc.add_heading("4.3 Protocole d'entraînement baseline", level=2)
    doc.add_paragraph(
        "Les deux modèles sont entraînés avec un seed fixé (42), le même dataset et un budget d'entraînement "
        "comparable : 1000 steps, batch size 128, sur la même machine CPU. Les checkpoints, logs de loss (CSV) "
        "et échantillons générés sont sauvegardés à intervalles réguliers (tous les 250 steps) via des scripts "
        "reproductibles (scripts/train.py --config <config>.yaml)."
    )
    add_table(
        doc,
        "Hyperparamètres des configurations baseline",
        ["Hyperparamètre", "DDPM (configs/ddpm_base.yaml)", "GAN (configs/gan_base.yaml)"],
        [
            ["Steps d'entraînement", "1000", "1000"],
            ["Batch size", "128", "128"],
            ["Seed", "42", "42"],
            ["Optimiseur", "Adam, lr = 2e-4", "Adam, lr = 2e-4, β1 = 0,5"],
            ["Architecture", "U-Net, 3,54M paramètres", "DCGAN (Generator + Discriminator)"],
            ["Spécifique", "T = 1000, schedule linéaire", "latent_dim = 100"],
        ],
    )


def build_ablation(doc: Document) -> None:
    doc.add_heading("5. Étude d'ablation — nombre de pas de diffusion", level=1)
    doc.add_paragraph(
        "Protocole : trois configurations DDPM identiques (même architecture, même seed 42, même budget "
        "d'entraînement de 500 steps) ne différant que par le nombre de pas de diffusion T (100, 400, 1000). "
        "Pour chaque configuration, le temps d'entraînement, le temps de génération (8 images) et la loss "
        "finale (moyenne des 20 derniers logs) sont mesurés."
    )
    add_table(
        doc,
        "Résultats de l'étude d'ablation (budget d'entraînement identique = 500 steps)",
        ["T (pas de diffusion)", "Temps entraînement", "Temps génération (8 img)", "Loss finale"],
        [
            ["100", "3295,7 s", "13,4 s", "0,1107"],
            ["400", "3667,6 s", "52,0 s", "0,0648"],
            ["1000", "3798,8 s", "129,0 s", "0,0404"],
        ],
    )
    add_figure(doc, "ablation_gentime_loss.png", "Temps de génération et loss finale en fonction de T")
    add_figure(doc, "ablation_samples.png", "Échantillons générés à budget d'entraînement identique (T=100, 400, 1000)")
    doc.add_paragraph(
        "Le temps de génération croît quasi linéairement avec T (ratio mesuré ≈ 1 / 3,9 / 9,6, proche du "
        "ratio théorique 1/4/10), confirmant le coût structurel du sampling multi-pas. La loss finale décroît "
        "avec T, mais mécaniquement : à T élevé, chaque pas ajoute moins de bruit, la tâche de débruitage par "
        "pas est plus facile — ce n'est pas directement une mesure de qualité perceptuelle."
    )
    doc.add_paragraph(
        "Résultat le plus notable, et contre-intuitif : à budget d'entraînement égal (500 steps), la "
        "configuration T=400 produit les échantillons visuellement les plus nets, tandis que T=1000 est "
        "visuellement moins net que T=400 malgré sa loss plus basse. Explication proposée : avec un nombre de "
        "steps de gradient fixe, chaque valeur de t est vue en moyenne moins souvent lorsque T est grand, donc "
        "le modèle est relativement sous-entraîné à T élevé. Ceci est cohérent avec le run baseline (T=1000, "
        "mais 1000 steps d'entraînement au lieu de 500), qui produisait des échantillons nettement plus nets — "
        "le nombre de steps d'entraînement nécessaire pour bien exploiter un T élevé semble donc croître avec T. "
        "Conclusion pratique : sur un budget de calcul CPU contraint, T=400 offre le meilleur compromis "
        "qualité/coût de génération parmi les configurations testées."
    )


def build_results(doc: Document) -> None:
    doc.add_heading("6. Résultats", level=1)

    doc.add_heading("6.1 Entraînement baseline — stabilité", level=2)
    add_table(
        doc,
        "Comparaison des entraînements baseline (1000 steps, budget comparable)",
        ["Métrique", "DDPM", "GAN"],
        [
            ["Durée totale", "7235,6 s (~2h00)", "1012,5 s (~17 min)"],
            ["Métrique finale", "Loss MSE : 1,044 → 0,041", "D(real) : 0,30 → 0,78 ; D(fake) : 0,33 → 0,24"],
            ["Comportement", "Décroissance lisse, monotone, sans oscillation", "Oscillant ; pics de loss_g jusqu'à 5,38 ; discriminateur prend l'avantage"],
            ["Mode collapse", "Non applicable", "Non observé (diversité de formes conservée)"],
        ],
    )
    add_figure(doc, "ddpm_baseline_samples.png", "Échantillons DDPM baseline (step 1000)")
    add_figure(doc, "gan_baseline_samples.png", "Échantillons GAN baseline (step 1000, artefacts en damier visibles)")
    doc.add_paragraph(
        "Le DDPM converge de façon prévisible et reproductible ; le GAN nécessite un suivi actif de l'équilibre "
        "générateur/discriminateur et reste plus instable, sans toutefois montrer de mode collapse total sur "
        "ce run (seed unique — cf. Discussion et limites)."
    )

    doc.add_heading("6.2 Comparaison chiffrée (FID, diversité, temps de génération)", level=2)
    doc.add_paragraph(
        "Comparaison sur un protocole identique : 100 images générées par modèle (à partir des checkpoints "
        "baseline), comparées à 100 images réelles du test set (notebooks/06_evaluation_ddpm_vs_gan.ipynb)."
    )
    add_table(
        doc,
        "Comparaison chiffrée DDPM vs GAN (100 images par groupe)",
        ["Métrique", "DDPM", "GAN"],
        [
            ["FID (vs réel)", "114,12", "173,47"],
            ["Variance intra-batch", "0,2443", "0,2292 (réel : 0,2840)"],
            ["Quasi-doublons (100 img, seuil 0,05)", "0", "0"],
            ["Temps de génération (100 images)", "1854,0 s (18,54 s/image)", "0,09 s (0,0009 s/image)"],
        ],
    )
    add_figure(doc, "eval_real_ddpm_gan.png", "Échantillons réels vs DDPM vs GAN (8 images par groupe)")
    doc.add_paragraph(
        "Le DDPM baseline obtient un FID nettement meilleur (114 contre 173), cohérent avec l'observation "
        "qualitative des échantillons (formes plus nettes, GAN affecté par des artefacts en damier). Ce score "
        "est calculé sur seulement 100 images par groupe (la littérature recommande plusieurs milliers pour un "
        "FID stable) et doit donc être interprété comme une comparaison relative sur un protocole identique, "
        "pas comme une valeur absolue comparable à la littérature. Le compromis inverse et structurel est le "
        "temps de génération : le GAN génère en une seule passe (~21 600 fois plus rapide que le DDPM sur ce "
        "protocole), le DDPM nécessitant T=1000 passes séquentielles du U-Net."
    )


def build_deployment(doc: Document) -> None:
    doc.add_heading("7. Déploiement", level=1)
    doc.add_paragraph(
        "Une démonstration applicative expose les deux modèles entraînés, conformément aux exigences du cours "
        "(niveau Master, exigence allégée — l'application illustre les résultats du rapport, pas un produit fini)."
    )
    doc.add_heading("7.1 API (FastAPI)", level=2)
    doc.add_paragraph(
        "app/api/main.py charge les deux modèles une fois au démarrage (chargement paresseux : si un checkpoint "
        "est absent, l'API démarre quand même). Deux endpoints : GET /health (statut, modèles chargés) et "
        "GET /generate?model=ddpm|gan&n=... (génère n images, retourne un JSON avec les images encodées en "
        "base64/PNG). Validation stricte du paramètre model (type Literal, réponse 422 si invalide), n plafonné "
        "par modèle (400 si dépassé, la génération DDPM étant coûteuse en CPU), 503 si le modèle demandé n'est "
        "pas chargé. Testée par 6 tests automatisés (fastapi.testclient.TestClient) et manuellement en conditions "
        "réelles (serveur uvicorn)."
    )
    doc.add_heading("7.2 Application web (Streamlit)", level=2)
    doc.add_paragraph(
        "app/web/app.py propose un sélecteur de modèle (DDPM / GAN / comparaison côte à côte), un bouton de "
        "génération appelant l'API (bibliothèque requests), l'affichage des images générées, et un bloc de "
        "métriques indicatives (FID de référence issu de l'évaluation, temps de génération mesuré en direct sur "
        "la requête en cours)."
    )
    doc.add_heading("7.3 Conteneurisation (Docker)", level=2)
    doc.add_paragraph(
        "Deux images Docker (python:3.11-slim), orchestrées par docker-compose.yml : le service api (port 8000) "
        "et le service web (port 8501, variable d'environnement API_URL réglée sur http://api:8000 pour la "
        "résolution DNS interne à Docker Compose). Les checkpoints entraînés (non versionnés dans le dépôt, trop "
        "volumineux) sont montés en volume en lecture seule plutôt que copiés dans l'image. Test de bout en bout "
        "réalisé (docker compose build puis up) : les deux conteneurs démarrent, /health confirme les modèles "
        "chargés depuis le volume, une génération GAN via l'API conteneurisée aboutit, la page Streamlit répond, "
        "et la connectivité réseau interne web → api a été vérifiée explicitement."
    )


def build_discussion(doc: Document) -> None:
    doc.add_heading("8. Discussion et limites", level=1)
    doc.add_paragraph(
        "Contrainte de calcul CPU : l'absence de GPU sur la machine de développement est le facteur limitant "
        "central de ce projet — il a conditionné le choix du dataset, la taille du U-Net (sans attention), le "
        "nombre de steps d'entraînement (1000, très inférieur aux dizaines de milliers habituelles pour un DDPM), "
        "et la taille des échantillons d'évaluation."
    )
    doc.add_paragraph(
        "FID sur échantillon réduit : le FID a été calculé sur seulement 100 images par groupe, faute de budget "
        "de calcul suffisant pour générer plusieurs milliers d'images DDPM (~18,5 s/image). Les valeurs absolues "
        "sont donc bruitées ; seule la comparaison relative DDPM vs GAN, sur un protocole identique, est robuste."
    )
    doc.add_paragraph(
        "Stabilité GAN non testée sur plusieurs seeds : la conclusion de stabilité (DDPM plus prévisible que "
        "GAN) repose sur un seul seed (42) par modèle. Une analyse rigoureuse de la stabilité d'entraînement "
        "nécessiterait plusieurs runs indépendants par modèle, ce qui n'a pas été réalisé faute de budget de "
        "calcul CPU (chaque run baseline supplémentaire coûtant environ 1 à 2 heures)."
    )
    doc.add_paragraph(
        "Étude d'ablation à budget réduit : les trois configurations d'ablation ont été entraînées sur 500 "
        "steps (contre 1000 pour la baseline), pour que les trois runs cumulés tiennent dans un temps de calcul "
        "raisonnable. Le résultat contre-intuitif observé (T=1000 visuellement moins net que T=400) est "
        "probablement en partie un artefact de ce budget réduit plutôt qu'une propriété générale du DDPM — une "
        "réplication à budget d'entraînement plus long serait nécessaire pour trancher définitivement."
    )
    doc.add_paragraph(
        "Répartition de l'implémentation : à la date de rédaction, l'ensemble des tâches Data et Model, ainsi "
        "qu'une partie substantielle des tâches Backend (API, app, Docker), ont été réalisées par un seul membre "
        "de l'équipe sur la branche de travail personnelle aaksp (cf. section Répartition du travail). "
        "L'intégration des contributions des deux autres membres de l'équipe reste à finaliser."
    )


def build_work_distribution(doc: Document) -> None:
    doc.add_heading("9. Répartition du travail", level=1)
    doc.add_paragraph(
        "Les rôles ont été attribués nommément (cf. Readme.md, CONTRIBUTING.md) : KOYE Leleda Ma Belle — "
        "Data / Experiment Engineer ; AHLI Kossi Sitsofe Pédro — Model / Research Engineer (principal), appui "
        "Reporting / Backend Developer ; KONTEVI Akossiwa Anne — Reporting / Backend Developer."
    )
    add_table(
        doc,
        "État d'avancement par domaine (source : HISTORY.md)",
        ["Domaine", "Réalisations", "Statut"],
        [
            ["Data", "Choix du dataset, EDA, pipeline de données, scripts de reproductibilité", "Complet"],
            ["Model", "Diffusion forward/inverse, U-Net, DCGAN, entraînements baseline, étude d'ablation, métriques (FID, diversité)", "Complet"],
            ["Backend — API/App/Docker", "FastAPI, Streamlit, conteneurisation Docker", "Complet"],
            ["Backend — Rapport", "Rédaction de ce document", "Premier jet"],
            ["Backend — Présentation", "Slides et démo live", "À faire"],
        ],
    )
    doc.add_paragraph(
        "À ce stade, l'ensemble des réalisations listées ci-dessus a été effectué par AHLI Kossi Sitsofe Pédro "
        "sur la branche de travail personnelle aaksp (non fusionnée sur master), en appui sur les deux autres "
        "rôles faute de contributions encore intégrées de KOYE Leleda Ma Belle et KONTEVI Akossiwa Anne. Cette "
        "section sera mise à jour pour refléter précisément la contribution de chaque membre une fois leurs "
        "apports intégrés au dépôt."
    )


def build_references(doc: Document) -> None:
    doc.add_heading("Références", level=1)
    refs = [
        "Goodfellow, I., Pouget-Abadie, J., Mirza, M., Xu, B., Warde-Farley, D., Ozair, S., Courville, A., & Bengio, Y. (2014). Generative Adversarial Networks. NeurIPS.",
        "Radford, A., Metz, L., & Chintala, S. (2015). Unsupervised Representation Learning with Deep Convolutional Generative Adversarial Networks. arXiv:1511.06434.",
        "Ho, J., Jain, A., & Abbeel, P. (2020). Denoising Diffusion Probabilistic Models. NeurIPS.",
        "Nichol, A., & Dhariwal, P. (2021). Improved Denoising Diffusion Probabilistic Models. ICML. arXiv:2102.09672.",
        "Odena, A., Dumoulin, V., & Olah, C. (2016). Deconvolution and Checkerboard Artifacts. Distill.",
        "Xiao, H., Rasul, K., & Vollgraf, R. (2017). Fashion-MNIST: a Novel Image Dataset for Benchmarking Machine Learning Algorithms. arXiv:1708.07747.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")


# ---------------------------------------------------------------------------
# Assemblage
# ---------------------------------------------------------------------------

def main() -> None:
    doc = Document()
    setup_styles(doc)

    build_cover(doc)
    build_sommaire(doc)
    build_toc(doc)
    build_glossary(doc)

    build_introduction(doc)
    build_related_work(doc)
    build_data(doc)
    build_method(doc)
    build_ablation(doc)
    build_results(doc)
    build_deployment(doc)
    build_discussion(doc)
    build_work_distribution(doc)
    build_references(doc)

    doc.save(OUTPUT_PATH)
    print(f"Rapport généré : {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
